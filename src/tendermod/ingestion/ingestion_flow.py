import logging
from glob import glob
from pathlib import Path

from tendermod.config.settings import CHROMA_PERSIST_DIR, ROOT_DIR
from tendermod.ingestion.pdf_loader import load_docs
from tendermod.ingestion.chunking import chunk_docs
from tendermod.retrieval.embeddings import embed_docs
from tendermod.retrieval.vectorstore import create_vectorstore

logger = logging.getLogger(__name__)


def export_ocr_to_docx(documents, pdf_filename: str) -> str:
    """Convierte Markdown estructurado (pymupdf4llm) a Word con headings y tablas."""
    from docx import Document as DocxDocument
    from docx.shared import Pt

    ocr_dir = Path(CHROMA_PERSIST_DIR).parent / "ocr"
    ocr_dir.mkdir(exist_ok=True)
    output_path = ocr_dir / (Path(pdf_filename).stem + "_ocr.docx")

    doc = DocxDocument()
    doc.add_heading("Documento OCR — " + Path(pdf_filename).name, level=1)

    for page_doc in documents:
        if not page_doc.page_content.strip():
            continue
        page_num = page_doc.metadata.get("page", "?")
        doc.add_heading(f"Página {int(page_num) + 1}", level=2)

        for line in page_doc.page_content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:].strip("*").strip(), level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:].strip("*").strip(), level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:].strip("*").strip(), level=1)
            elif stripped.startswith("|"):
                # Fila de tabla Markdown → monospace para preservar alineación
                p = doc.add_paragraph(stripped)
                if p.runs:
                    p.runs[0].font.name = "Courier New"
                    p.runs[0].font.size = Pt(8)
            elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
                p = doc.add_paragraph()
                p.add_run(stripped[2:-2]).bold = True
            else:
                doc.add_paragraph(stripped)

    doc.save(str(output_path))
    logger.info("[export_ocr_to_docx] Guardado: %s", output_path)
    return str(output_path)


def ingest_documents() -> dict:
    """
    Ingesta el PDF en ChromaDB.
    Aplica OCR automáticamente si el PDF es escaneado.
    Retorna dict: {vectorstore, ocr_applied, ocr_docx_path}.
    """
    logger.info("[ingest_documents] Iniciando ingesta")

    docs, ocr_applied = load_docs()

    chunks = chunk_docs(docs)

    # Filtro defensivo: eliminar chunks con contenido vacío
    chunks_validos = [c for c in chunks if c.page_content.strip()]
    descartados = len(chunks) - len(chunks_validos)
    if descartados:
        logger.warning(
            "[ingest_documents] %d chunks vacíos descartados de %d totales",
            descartados, len(chunks),
        )

    if not chunks_validos:
        raise ValueError(
            "El PDF no contiene texto extraíble. "
            "Verifique que Tesseract esté instalado para PDFs escaneados."
        )

    logger.info("[ingest_documents] %d chunks válidos → embeddings + ChromaDB", len(chunks_validos))
    vectorstore = create_vectorstore(chunks_validos, embed_docs(), path=CHROMA_PERSIST_DIR)

    ocr_docx_path = None
    if ocr_applied:
        pdf_files = glob(str(ROOT_DIR / "data" / "*.pdf"))
        if pdf_files:
            try:
                ocr_docx_path = export_ocr_to_docx(docs, pdf_files[0])
            except Exception as exc:
                logger.error("[ingest_documents] Error generando docx OCR: %s", exc)

    return {
        "vectorstore": vectorstore,
        "ocr_applied": ocr_applied,
        "ocr_docx_path": ocr_docx_path,
    }
